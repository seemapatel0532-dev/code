# app/services/resume_builder_service.py
from __future__ import annotations

from io import BytesIO
from typing import Any, Dict, List

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT

from reportlab.lib import colors
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import (
    HRFlowable,
    ListFlowable,
    ListItem,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)
from reportlab.platypus.doctemplate import BaseDocTemplate, Frame, PageTemplate
from reportlab.platypus.flowables import Flowable


# -----------------------------
# Shared helpers
# -----------------------------
def _get(profile: Dict[str, Any], key: str, default=None):
    return profile.get(key) if isinstance(profile, dict) else default


def _safe_str(x) -> str:
    return (x or "").strip()


def _norm_url(u: str) -> str:
    u = _safe_str(u)
    if not u:
        return ""
    if u.startswith("http://") or u.startswith("https://"):
        return u
    if u.startswith("www."):
        return "https://" + u
    if "." in u and " " not in u:
        return "https://" + u
    return u


def _mailto(email: str) -> str:
    email = _safe_str(email)
    return f"mailto:{email}" if email else ""


def _tel(phone: str) -> str:
    phone = _safe_str(phone)
    if not phone:
        return ""
    compact = "".join(ch for ch in phone if ch.isdigit() or ch == "+")
    return f"tel:{compact}" if compact else ""


def _split_links_field(links: str) -> List[str]:
    links = _safe_str(links)
    if not links:
        return []
    parts: List[str] = []
    for chunk in links.replace("\n", " ").split("|"):
        chunk = chunk.strip()
        if not chunk:
            continue
        for p in chunk.split(","):
            p = p.strip()
            if p:
                parts.append(p)

    seen = set()
    out: List[str] = []
    for p in parts:
        k = p.lower()
        if k not in seen:
            seen.add(k)
            out.append(p)
    return out[:6]


# ============================================================
# DOCX (structured + "interactive" via collapsible headings + links)
# ============================================================
def _set_doc_defaults(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)


def _add_hyperlink(paragraph, url: str, text: str) -> None:
    url = _safe_str(url)
    text = _safe_str(text) or url
    if not url:
        paragraph.add_run(text)
        return

    url = _norm_url(url)

    part = paragraph.part
    r_id = part.relate_to(
        url,
        reltype="http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    c = OxmlElement("w:color")
    c.set(qn("w:val"), "1E3A8A")
    rPr.append(c)

    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)

    new_run.append(rPr)

    t = OxmlElement("w:t")
    t.text = text
    new_run.append(t)

    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def _heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)


def _add_divider(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run("─" * 50)
    run.font.size = Pt(8)


def build_docx(profile: Dict[str, Any]) -> bytes:
    doc = Document()
    _set_doc_defaults(doc)

    personal = _get(profile, "personal", {}) or {}
    name = _safe_str(personal.get("name")) or "Your Name"
    email = _safe_str(personal.get("email"))
    phone = _safe_str(personal.get("phone"))
    links_field = _safe_str(personal.get("links"))
    links = _split_links_field(links_field)

    # Title
    title = doc.add_paragraph()
    title_run = title.add_run(name)
    title_run.bold = True
    title_run.font.size = Pt(24)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title.paragraph_format.space_after = Pt(2)

    # Meta row as table (aligned)
    meta_tbl = doc.add_table(rows=1, cols=3)
    meta_tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    meta_tbl.autofit = True
    for c in meta_tbl.rows[0].cells:
        c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    # Email
    p0 = meta_tbl.cell(0, 0).paragraphs[0]
    if email:
        _add_hyperlink(p0, _mailto(email), email)

    # Phone
    p1 = meta_tbl.cell(0, 1).paragraphs[0]
    if phone:
        _add_hyperlink(p1, _tel(phone), phone)

    # Links
    p2 = meta_tbl.cell(0, 2).paragraphs[0]
    if links:
        for i, u in enumerate(links):
            if i > 0:
                p2.add_run("  •  ")
            uu = _norm_url(u)
            label = u.replace("https://", "").replace("http://", "")
            _add_hyperlink(p2, uu, label)
    elif links_field:
        p2.add_run(links_field)

    _add_divider(doc)

    # Summary
    summary = _safe_str(profile.get("summary"))
    if summary:
        _heading(doc, "Summary", 1)
        p = doc.add_paragraph(summary)
        p.paragraph_format.space_after = Pt(4)

    # Skills (structured 3-col grid)
    skills = [s for s in (profile.get("skills") or []) if _safe_str(s)]
    if skills:
        _heading(doc, "Skills", 1)
        cols = 3
        rows = (len(skills) + cols - 1) // cols
        tbl = doc.add_table(rows=rows, cols=cols)
        tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
        tbl.autofit = True

        idx = 0
        for r in range(rows):
            for c in range(cols):
                cell = tbl.cell(r, c)
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                cell.text = ""
                if idx < len(skills):
                    run = cell.paragraphs[0].add_run(skills[idx])
                    run.bold = True
                    idx += 1
        doc.add_paragraph()

    # Education
    edu = profile.get("education") or []
    if edu:
        _heading(doc, "Education", 1)
        for e in edu:
            degree = _safe_str(e.get("degree"))
            inst = _safe_str(e.get("institution"))
            year = _safe_str(e.get("year"))
            line = " • ".join([x for x in [degree, inst, year] if x])
            p = doc.add_paragraph(line)
            p.paragraph_format.space_after = Pt(2)

    # Experience
    exp = profile.get("experience") or []
    if exp:
        _heading(doc, "Experience", 1)
        for job in exp:
            role = _safe_str(job.get("role"))
            company = _safe_str(job.get("company"))
            duration = _safe_str(job.get("duration"))

            head = doc.add_paragraph()
            r0 = head.add_run(" • ".join([x for x in [role, company] if x]) or "Experience")
            r0.bold = True
            r0.font.size = Pt(11.5)
            if duration:
                head.add_run(f"  ({duration})").italic = True

            for b in (job.get("bullets") or []):
                b = _safe_str(b)
                if b:
                    doc.add_paragraph(b, style="List Bullet")

    # Projects (collapsible per project in Word)
    proj = profile.get("projects") or []
    if proj:
        _heading(doc, "Projects", 1)
        for p in proj:
            pname = _safe_str(p.get("name")) or "Project"
            tech = _safe_str(p.get("tech"))

            _heading(doc, pname, 2)
            if tech:
                tech_p = doc.add_paragraph()
                tech_run = tech_p.add_run(f"Tech: {tech}")
                tech_run.italic = True
                tech_p.paragraph_format.space_after = Pt(2)

            for b in (p.get("bullets") or []):
                b = _safe_str(b)
                if b:
                    doc.add_paragraph(b, style="List Bullet")

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ============================================================
# PDF (structured + attractive + bookmarks + clickable links)
# ============================================================
class _SectionMarker(Flowable):
    def __init__(self, key: str, title: str):
        super().__init__()
        self.key = key
        self.title = title

    def draw(self):
        pass


class _ResumeDocTemplate(BaseDocTemplate):
    def __init__(self, filename_or_buf, **kwargs):
        super().__init__(filename_or_buf, **kwargs)
        frame = Frame(self.leftMargin, self.bottomMargin, self.width, self.height, id="normal")
        self.addPageTemplates([PageTemplate(id="Resume", frames=[frame], onPage=self._on_page)])

    def _on_page(self, canvas, doc):
        canvas.saveState()
        canvas.setFont("Helvetica", 9)
        canvas.setFillColor(colors.grey)
        canvas.drawRightString(doc.pagesize[0] - doc.rightMargin, 0.55 * inch, f"Page {doc.page}")
        canvas.restoreState()

    def afterFlowable(self, flowable):
        if isinstance(flowable, _SectionMarker):
            key = flowable.key
            title = flowable.title
            c = self.canv
            c.bookmarkPage(key)
            c.addOutlineEntry(title, key, level=0, closed=False)


def _pdf_styles():
    base = getSampleStyleSheet()

    title = ParagraphStyle(
        "Title",
        parent=base["Title"],
        fontName="Helvetica-Bold",
        fontSize=22,
        leading=26,
        textColor=colors.HexColor("#0f172a"),
        spaceAfter=6,
    )
    meta = ParagraphStyle(
        "Meta",
        parent=base["Normal"],
        fontName="Helvetica",
        fontSize=10.5,
        leading=14,
        textColor=colors.HexColor("#334155"),
        spaceAfter=10,
    )
    summary = ParagraphStyle(
        "SummaryText",
        parent=base["Normal"],
        fontName="Helvetica",
        fontSize=10.5,
        leading=15,
        textColor=colors.HexColor("#0f172a"),
        spaceAfter=8,
    )
    h = ParagraphStyle(
        "H1",
        parent=base["Heading2"],
        fontName="Helvetica-Bold",
        fontSize=12.5,
        leading=16,
        textColor=colors.HexColor("#1E3A8A"),
        spaceBefore=10,
        spaceAfter=6,
    )
    item_title = ParagraphStyle(
        "ItemTitle",
        parent=base["Normal"],
        fontName="Helvetica-Bold",
        fontSize=10.8,
        leading=14,
        textColor=colors.HexColor("#0f172a"),
        spaceAfter=2,
    )
    item_sub = ParagraphStyle(
        "ItemSub",
        parent=base["Normal"],
        fontName="Helvetica-Oblique",
        fontSize=9.8,
        leading=13,
        textColor=colors.HexColor("#475569"),
        spaceAfter=4,
    )
    bullet = ParagraphStyle(
        "Bullet",
        parent=base["Normal"],
        fontName="Helvetica",
        fontSize=10.2,
        leading=14,
        leftIndent=14,
        bulletIndent=6,
        textColor=colors.HexColor("#0f172a"),
        spaceAfter=2,
    )
    # skills chip text (wrap-friendly)
    chip = ParagraphStyle(
        "Chip",
        parent=base["Normal"],
        fontName="Helvetica",
        fontSize=10,
        leading=12.5,
        textColor=colors.HexColor("#0f172a"),
    )

    return {"title": title, "meta": meta, "summary": summary, "h": h, "item_title": item_title, "item_sub": item_sub, "bullet": bullet, "chip": chip}


def _skills_grid(skills: List[str], cols: int = 3) -> List[List[str]]:
    """
    Fixed-column grid so it never overflows horizontally.
    """
    cols = max(2, min(4, int(cols)))
    skills = [s for s in skills if _safe_str(s)]
    rows = (len(skills) + cols - 1) // cols
    data: List[List[str]] = []
    idx = 0
    for _ in range(rows):
        row: List[str] = []
        for _c in range(cols):
            row.append(skills[idx] if idx < len(skills) else "")
            idx += 1
        data.append(row)
    return data


def build_pdf(profile: Dict[str, Any]) -> bytes:
    buf = BytesIO()

    doc = _ResumeDocTemplate(
        buf,
        pagesize=LETTER,
        leftMargin=0.7 * inch,
        rightMargin=0.7 * inch,
        topMargin=0.7 * inch,
        bottomMargin=0.65 * inch,
        title="CareerSetu Resume",
        author="CareerSetu AI",
    )

    st = _pdf_styles()
    story: List[Any] = []

    personal = _get(profile, "personal", {}) or {}
    name = _safe_str(personal.get("name")) or "Your Name"
    email = _safe_str(personal.get("email"))
    phone = _safe_str(personal.get("phone"))
    links_field = _safe_str(personal.get("links"))
    links = _split_links_field(links_field)

    # Header
    story.append(Paragraph(name, st["title"]))

    meta_parts = []
    if email:
        meta_parts.append(f'<a href="{_mailto(email)}" color="#1E3A8A">{email}</a>')
    if phone:
        meta_parts.append(f'<a href="{_tel(phone)}" color="#1E3A8A">{phone}</a>')
    for u in links:
        nu = _norm_url(u)
        label = u.replace("https://", "").replace("http://", "")
        meta_parts.append(f'<a href="{nu}" color="#1E3A8A">{label}</a>')

    if meta_parts:
        story.append(Paragraph(" &nbsp;&nbsp;•&nbsp;&nbsp; ".join(meta_parts), st["meta"]))
    else:
        story.append(Spacer(1, 6))

    story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor("#e2e8f0"), spaceBefore=2, spaceAfter=10))

    # Summary
    summary = _safe_str(profile.get("summary"))
    if summary:
        story.append(_SectionMarker("sec_summary", "Summary"))
        story.append(Paragraph("Summary", st["h"]))
        story.append(Paragraph(summary, st["summary"]))

    # Skills (FIXED GRID + WRAP + FIXED COL WIDTHS) ✅
    skills = [s for s in (profile.get("skills") or []) if _safe_str(s)]
    if skills:
        story.append(_SectionMarker("sec_skills", "Skills"))
        story.append(Paragraph("Skills", st["h"]))

        cols = 3
        data_raw = _skills_grid(skills, cols=cols)

        # Put Paragraph objects in each cell so long skills wrap inside the cell
        data = []
        for row in data_raw:
            data.append([Paragraph(_safe_str(cell), st["chip"]) if _safe_str(cell) else "" for cell in row])

        available_w = doc.width  # exact usable width on page
        col_w = available_w / cols
        col_widths = [col_w] * cols

        tbl = Table(data, colWidths=col_widths, hAlign="LEFT")
        tbl.setStyle(
            TableStyle(
                [
                    ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                    ("LEFTPADDING", (0, 0), (-1, -1), 8),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 8),
                    ("TOPPADDING", (0, 0), (-1, -1), 6),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
                    ("GRID", (0, 0), (-1, -1), 0.6, colors.HexColor("#e2e8f0")),
                    ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#F8FAFC")),
                ]
            )
        )
        story.append(tbl)
        story.append(Spacer(1, 8))

    # Education
    edu = profile.get("education") or []
    if edu:
        story.append(_SectionMarker("sec_education", "Education"))
        story.append(Paragraph("Education", st["h"]))
        for e in edu:
            degree = _safe_str(e.get("degree"))
            inst = _safe_str(e.get("institution"))
            year = _safe_str(e.get("year"))
            line = " • ".join([x for x in [degree, inst] if x])
            story.append(Paragraph(line or "Education", st["item_title"]))
            if year:
                story.append(Paragraph(year, st["item_sub"]))
        story.append(Spacer(1, 4))

    # Experience
    exp = profile.get("experience") or []
    if exp:
        story.append(_SectionMarker("sec_experience", "Experience"))
        story.append(Paragraph("Experience", st["h"]))
        for job in exp:
            role = _safe_str(job.get("role"))
            company = _safe_str(job.get("company"))
            duration = _safe_str(job.get("duration"))

            title_line = " • ".join([x for x in [role, company] if x]) or "Experience"
            story.append(Paragraph(title_line, st["item_title"]))
            if duration:
                story.append(Paragraph(duration, st["item_sub"]))

            bullets = [b for b in (job.get("bullets") or []) if _safe_str(b)]
            if bullets:
                lf = ListFlowable(
                    [ListItem(Paragraph(_safe_str(b), st["bullet"])) for b in bullets],
                    bulletType="bullet",
                    start="•",
                    leftIndent=14,
                )
                story.append(lf)
            story.append(Spacer(1, 6))

    # Projects
    proj = profile.get("projects") or []
    if proj:
        story.append(_SectionMarker("sec_projects", "Projects"))
        story.append(Paragraph("Projects", st["h"]))

        for p in proj:
            pname = _safe_str(p.get("name")) or "Project"
            tech = _safe_str(p.get("tech"))
            story.append(Paragraph(pname, st["item_title"]))
            if tech:
                story.append(Paragraph(f"Tech: {tech}", st["item_sub"]))

            bullets = [b for b in (p.get("bullets") or []) if _safe_str(b)]
            if bullets:
                lf = ListFlowable(
                    [ListItem(Paragraph(_safe_str(b), st["bullet"])) for b in bullets],
                    bulletType="bullet",
                    start="•",
                    leftIndent=14,
                )
                story.append(lf)
            story.append(Spacer(1, 6))

    if not story:
        story.append(Paragraph("No resume content found. Please edit your profile in Builder.", st["summary"]))

    doc.build(story)
    return buf.getvalue()