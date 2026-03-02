from __future__ import annotations

from io import BytesIO
from typing import Any, Dict

from docx import Document
from docx.shared import Inches, Pt

from reportlab.lib import colors
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import Paragraph, SimpleDocTemplate, HRFlowable


def _safe_str(x: Any) -> str:
    return (x or "").strip()


def build_cover_letter_docx(profile: Dict[str, Any], letter_text: str) -> bytes:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    personal = (profile or {}).get("personal", {}) or {}
    name = _safe_str(personal.get("name")) or "Your Name"
    email = _safe_str(personal.get("email"))
    phone = _safe_str(personal.get("phone"))

    title = doc.add_paragraph()
    run = title.add_run(name)
    run.bold = True
    run.font.size = Pt(20)

    if email or phone:
        meta = doc.add_paragraph(" | ".join([x for x in [email, phone] if x]))
        meta.paragraph_format.space_after = Pt(8)

    doc.add_paragraph("")

    for block in (letter_text or "").split("\n\n"):
        p = doc.add_paragraph(block.strip())
        p.paragraph_format.space_after = Pt(8)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def build_cover_letter_pdf(profile: Dict[str, Any], letter_text: str) -> bytes:
    buf = BytesIO()
    doc = SimpleDocTemplate(
        buf,
        pagesize=LETTER,
        leftMargin=0.85 * inch,
        rightMargin=0.85 * inch,
        topMargin=0.85 * inch,
        bottomMargin=0.75 * inch,
        title="CareerSetu Cover Letter",
        author="CareerSetu AI",
    )

    styles = getSampleStyleSheet()
    title = ParagraphStyle(
        "CLTitle",
        parent=styles["Title"],
        fontName="Helvetica-Bold",
        fontSize=18,
        leading=22,
        textColor=colors.HexColor("#0f172a"),
        spaceAfter=6,
    )
    meta = ParagraphStyle(
        "CLMeta",
        parent=styles["Normal"],
        fontName="Helvetica",
        fontSize=10.5,
        leading=14,
        textColor=colors.HexColor("#334155"),
        spaceAfter=10,
    )
    body = ParagraphStyle(
        "CLBody",
        parent=styles["Normal"],
        fontName="Helvetica",
        fontSize=11,
        leading=16,
        textColor=colors.HexColor("#0f172a"),
        spaceAfter=10,
    )

    personal = (profile or {}).get("personal", {}) or {}
    name = _safe_str(personal.get("name")) or "Your Name"
    email = _safe_str(personal.get("email"))
    phone = _safe_str(personal.get("phone"))

    story = []
    story.append(Paragraph(name, title))
    if email or phone:
        story.append(Paragraph(" • ".join([x for x in [email, phone] if x]), meta))

    story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor("#e2e8f0"), spaceBefore=2, spaceAfter=12))

    text = _safe_str(letter_text) or "(No cover letter content)"
    for block in text.split("\n\n"):
        block = block.strip()
        if block:
            story.append(Paragraph(block.replace("\n", "<br/>"), body))

    doc.build(story)
    return buf.getvalue()